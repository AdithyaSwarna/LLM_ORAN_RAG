import os
import re
import fitz  # PyMuPDF
import json
import logging
from docx import Document

# Setup logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

def clean_text(text):
    """
    Clean text by removing unwanted lines, excessive whitespace, and specific patterns.
    """
    lines = text.splitlines()  # Split text into individual lines
    cleaned_lines = [
        line.strip()
        for line in lines
        if line.strip() and not re.match(r"_{5,}|-{5,}|\.{5,}", line)  # Remove lines with repetitive patterns
    ]
    cleaned_text = " ".join(cleaned_lines)  # Join lines into a single compact string
    cleaned_text = re.sub(r"\.{5,}", "", cleaned_text)
    return cleaned_text

def extract_pdf_metadata_text(pdf_path):
    """
    Extract metadata and text from a PDF file.
    """
    with fitz.open(pdf_path) as doc:
        metadata = doc.metadata or {}
        metadata["File Name"] = os.path.basename(pdf_path)
        metadata["PageCount"] = len(doc)

        text_content = {}
        for i, page in enumerate(doc):
            raw_text = page.get_text("text")
            cleaned_text_content = clean_text(raw_text)
            text_content[f"Page_{i + 1}"] = cleaned_text_content
    return metadata, text_content

def extract_docx_metadata_text(docx_path):
    """
    Extract metadata and text from a Word document.
    """
    doc = Document(docx_path)
    metadata = {
        "File Name": os.path.basename(docx_path),
        "PageCount": len(doc.paragraphs)
    }
    text_content = {}
    for i, paragraph in enumerate(doc.paragraphs):
        raw_text = paragraph.text
        cleaned_text_content = clean_text(raw_text)
        text_content[f"Paragraph_{i + 1}"] = cleaned_text_content
    return metadata, text_content

def save_json(data, output_path):
    """
    Save data to a JSON file.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def process_document(file_path, output_dir):
    """
    Process a document (PDF or Word) to extract metadata and text.
    """
    base_name, ext = os.path.splitext(os.path.basename(file_path))
    try:
        if ext.lower() == ".pdf":
            metadata, text_content = extract_pdf_metadata_text(file_path)
        elif ext.lower() == ".docx":
            metadata, text_content = extract_docx_metadata_text(file_path)
        else:
            logging.warning(f"Unsupported format for file: {file_path}")
            return

        # Save metadata and text
        os.makedirs(output_dir, exist_ok=True)
        save_json(metadata, os.path.join(output_dir, f"{base_name}_metadata.json"))
        save_json(text_content, os.path.join(output_dir, f"{base_name}_text.json"))
        logging.info(f"Processed: {file_path}")

    except Exception as e:
        logging.error(f"Error processing {file_path}: {str(e)}")

def main():
    input_dir = "/home/sswarna/Documents/oran_docs"
    output_base_dir = "/home/sswarna/Documents/oran_docs/output_all"

    for year in ["2022", "2023", "2024"]:
        year_input_dir = os.path.join(input_dir, year)
        year_output_dir = os.path.join(output_base_dir, f"Output_{year}")

        if not os.path.exists(year_input_dir):
            logging.warning(f"Input directory for {year} does not exist.")
            continue

        for file_name in os.listdir(year_input_dir):
            file_path = os.path.join(year_input_dir, file_name)
            process_document(file_path, year_output_dir)

if __name__ == "__main__":
    main()
